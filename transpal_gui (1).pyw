# -*- coding: utf-8 -*-
"""
TransPal — Trợ lý dịch thuật (bản GUI, tiếng Việt có dấu, chạy bằng double-click)
=================================================================================
Cài đặt (chỉ 1 lần, mở CMD):
  pip install keyboard pyperclip pyautogui google-genai deep-translator python-docx pypdf

Đặt API key (chỉ 1 lần — cách bền vững, không cần gõ lại):
  setx GEMINI_API_KEY "key_cua_ban"
  (đóng mở lại ứng dụng sau khi setx)

Chạy: double-click vào file transpal_gui.pyw này là xong, không cần terminal.
"""

import os
import sys
import json
import time
import queue
import ctypes
import threading
import traceback

import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox

# --- import thư viện ngoài, nếu thiếu thì báo bằng hộp thoại (vì không có console) ---
try:
    import keyboard
    import pyperclip
    import pyautogui
except ImportError as _e:
    _root = tk.Tk(); _root.withdraw()
    messagebox.showerror(
        "TransPal — thiếu thư viện",
        f"Thiếu thư viện: {_e.name}\n\nMở CMD và chạy:\n"
        "pip install keyboard pyperclip pyautogui google-genai deep-translator python-docx pypdf")
    sys.exit(1)

# ============================================================
# 1. CẤU HÌNH
# ============================================================
# Luôn làm việc trong thư mục chứa file script (double-click đôi khi
# khởi động ở thư mục khác, khiến không tìm thấy context.md / config.json)
APP_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(APP_DIR, "config.json")
CONTEXT_FILE = os.path.join(APP_DIR, "context.md")

DEFAULT_CONFIG = {
    "hotkey": "alt+shift+t",
    "quit_hotkey": "ctrl+alt+x",
    "toggle_mode_hotkey": "ctrl+alt+m",
    "mode": "ai",                # "ai" | "basic"
    "target_lang": "vi",
    "auto_replace": False,
    "restore_clipboard": True,
    "models": [
        "gemma-4-26b-a4b-it",
        "gemma-4-31b-it",
        "gemini-3.1-flash-lite",
        "gemini-2.5-flash"
    ]
}

DEFAULT_CONTEXT = (
    "# Hướng dẫn dịch thuật\n"
    "- Dịch tự nhiên, mượt mà, hợp văn phong Việt Nam.\n"
    "- Giữ nguyên thuật ngữ kỹ thuật chuyên ngành nếu không cần dịch.\n"
    "- Giữ nguyên format gốc (xuống dòng, danh sách, code...).\n"
    "- Tuyệt đối không thêm lời giải thích của AI vào kết quả.\n"
)


def load_config():
    if not os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(DEFAULT_CONFIG, f, ensure_ascii=False, indent=2)
        return dict(DEFAULT_CONFIG)
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        cfg = json.load(f)
    for k, v in DEFAULT_CONFIG.items():
        cfg.setdefault(k, v)
    # --- Tự sửa tên model sai / nâng cấp danh sách cũ trong config.json sẵn có ---
    name_fixes = {
        "gemma-4-26b": "gemma-4-26b-a4b-it",   # thiếu hậu tố -it là API báo 404
        "gemma-4-31b": "gemma-4-31b-it",
    }
    cfg["models"] = [name_fixes.get(m, m) for m in cfg.get("models", [])]
    old_defaults = {"gemini-2.5-flash", "gemini-2.5-flash-lite", "gemini-2.0-flash"}
    if not cfg["models"] or set(cfg["models"]) <= old_defaults:
        cfg["models"] = list(DEFAULT_CONFIG["models"])
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)
    return cfg


def save_config(cfg):
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


CONFIG = load_config()


def read_context_file():
    if not os.path.exists(CONTEXT_FILE):
        with open(CONTEXT_FILE, "w", encoding="utf-8") as f:
            f.write(DEFAULT_CONTEXT)
    with open(CONTEXT_FILE, "r", encoding="utf-8") as f:
        return f.read().strip()


def write_context_file(content):
    with open(CONTEXT_FILE, "w", encoding="utf-8") as f:
        f.write(content.strip() + "\n")


# ============================================================
# 2. ĐỌC FILE IMPORT (.txt .md .docx .pdf)
# ============================================================
def extract_file_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md"):
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            raise RuntimeError("Thiếu thư viện python-docx. Mở CMD chạy: pip install python-docx")
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("Thiếu thư viện pypdf. Mở CMD chạy: pip install pypdf")
        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    raise RuntimeError(f"Chưa hỗ trợ định dạng {ext}. Hỗ trợ: .txt .md .docx .pdf")


# ============================================================
# 3. AI (Gemini) + DỊCH CƠ BẢN
# ============================================================
_genai_client = None


def get_genai_client():
    global _genai_client
    if _genai_client is None:
        api_key = os.environ.get("GEMINI_API_KEY", "").strip()
        if not api_key:
            raise RuntimeError("Chưa đặt biến môi trường GEMINI_API_KEY.")
        from google import genai
        _genai_client = genai.Client(api_key=api_key)
    return _genai_client


def call_gemini(prompt, log=print):
    """Gọi Gemini, tự fallback qua các model. Trả về (text, tên_model)."""
    client = get_genai_client()
    last_err = None
    for model_name in CONFIG["models"]:
        try:
            log(f"   … đang gọi model: {model_name}")
            resp = client.models.generate_content(model=model_name, contents=prompt)
            text = (resp.text or "").strip()
            if text:
                return text, model_name
        except Exception as e:
            last_err = e
            msg = str(e)
            if "429" in msg or "RESOURCE_EXHAUSTED" in msg:
                log(f"   ⚠ {model_name}: hết quota (429) → chuyển model kế tiếp")
            elif "404" in msg or "NOT_FOUND" in msg:
                log(f"   ⚠ {model_name}: không tìm thấy model (404, sai tên?) → chuyển model kế tiếp")
            else:
                log(f"   ⚠ model {model_name} lỗi: {msg}")
    raise RuntimeError(f"Tất cả model AI đều thất bại. Lỗi cuối: {last_err}")


def strip_md_fence(text):
    """Bỏ ```markdown … ``` nếu AI lỡ bọc kết quả trong code fence."""
    t = text.strip()
    if t.startswith("```"):
        lines = t.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]
        t = "\n".join(lines).strip()
    return t


def ai_convert_and_merge(raw_text, current_context, log=print):
    """AI chuyển nội dung file import sang Markdown chuẩn + merge với context cũ."""
    prompt = (
        "Bạn là trợ lý biên tập tài liệu hướng dẫn dịch thuật.\n\n"
        "## NHIỆM VỤ\n"
        "1. Đọc NỘI DUNG MỚI được import bên dưới (có thể là text thô, trích từ PDF/Word, "
        "lộn xộn, thiếu cấu trúc).\n"
        "2. Chuyển nó thành Markdown gọn gàng: tiêu đề (#, ##), gạch đầu dòng, "
        "bảng thuật ngữ nếu có các cặp từ tương ứng.\n"
        "3. MERGE với FILE CONTEXT HIỆN TẠI: giữ lại toàn bộ quy tắc cũ còn giá trị, "
        "gộp mục trùng lặp, nếu mâu thuẫn thì ưu tiên NỘI DUNG MỚI, sắp xếp lại theo "
        "nhóm hợp lý (ví dụ: Quy tắc chung / Thuật ngữ / Văn phong / Lưu ý khác).\n"
        "4. Kết quả phải là một file hướng dẫn dịch thuật hoàn chỉnh, ngắn gọn, "
        "dễ cho AI dịch thuật khác đọc hiểu.\n\n"
        "## RÀNG BUỘC\n"
        "- CHỈ trả về nội dung Markdown của file kết quả, không lời dẫn, "
        "không bọc trong ```.\n"
        "- Viết bằng tiếng Việt (giữ nguyên thuật ngữ gốc khi cần).\n\n"
        f"## FILE CONTEXT HIỆN TẠI\n{current_context}\n\n"
        f"## NỘI DUNG MỚI ĐƯỢC IMPORT\n{raw_text}"
    )
    result, model = call_gemini(prompt, log=log)
    return strip_md_fence(result), model


def translate_basic(text, target_lang):
    from deep_translator import GoogleTranslator
    return GoogleTranslator(source="auto", target=target_lang).translate(text)


def translate_ai(text, target_lang, log=print):
    context_content = read_context_file()
    prompt = (
        f"Bạn là một biên dịch viên chuyên nghiệp. Hướng dẫn/ngữ cảnh:\n"
        f"{context_content}\n\n"
        f"Hãy dịch đoạn văn dưới đây sang ngôn ngữ có mã '{target_lang}'. "
        f"Giữ nguyên format gốc. CHỈ trả về bản dịch, không giải thích gì thêm.\n\n"
        f"--- VĂN BẢN GỐC ---\n{text}"
    )
    return call_gemini(prompt, log=log)


def do_translate(text, log=print):
    mode = CONFIG["mode"]
    lang = CONFIG["target_lang"]
    if mode == "basic":
        return translate_basic(text, lang), "Google Translate"
    try:
        result, model = translate_ai(text, lang, log=log)
        return result, f"AI ({model})"
    except Exception as e:
        log(f"⚠ AI thất bại ({e}) → chuyển sang dịch cơ bản…")
        return translate_basic(text, lang), "Google Translate (dự phòng)"


# ============================================================
# 4. CLIPBOARD & PHÍM
# ============================================================
def wait_hotkey_released(hotkey, timeout=5.0):
    """Chờ người dùng nhả hết các phím trong tổ hợp trước khi giả lập Ctrl+C."""
    keys = [k.strip() for k in hotkey.split("+")]
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            if not any(keyboard.is_pressed(k) for k in keys):
                time.sleep(0.05)
                return True
        except Exception:
            return True
        time.sleep(0.03)
    return False


def get_selected_text():
    """Xóa clipboard → Ctrl+C → poll đến khi có nội dung mới (chống race condition)."""
    try:
        old_clipboard = pyperclip.paste()
    except Exception:
        old_clipboard = ""
    pyperclip.copy("")
    time.sleep(0.05)
    keyboard.press_and_release("ctrl+c")
    for _ in range(48):
        time.sleep(0.025)
        try:
            txt = pyperclip.paste()
        except Exception:
            txt = ""
        if txt:
            return txt, old_clipboard
    return "", old_clipboard


def paste_text(text, target_hwnd=None):
    """Trả focus về cửa sổ gốc rồi dán đè bản dịch vào vùng đang bôi đen."""
    if target_hwnd and sys.platform == "win32":
        try:
            ctypes.windll.user32.SetForegroundWindow(target_hwnd)
            time.sleep(0.25)
        except Exception:
            pass
    old = None
    if CONFIG.get("restore_clipboard"):
        try:
            old = pyperclip.paste()
        except Exception:
            old = None
    pyperclip.copy(text)
    time.sleep(0.12)
    keyboard.press_and_release("ctrl+v")
    if CONFIG.get("restore_clipboard") and old is not None:
        def _restore():
            time.sleep(0.6)
            try:
                pyperclip.copy(old)
            except Exception:
                pass
        threading.Thread(target=_restore, daemon=True).start()


def get_foreground_hwnd():
    if sys.platform == "win32":
        try:
            return ctypes.windll.user32.GetForegroundWindow()
        except Exception:
            return None
    return None


# ============================================================
# 5. GIAO DIỆN CHÍNH
# ============================================================
BG = "#1e1e2e"
BG2 = "#2a2a3c"
FG = "#cdd6f4"
FG_DIM = "#a6adc8"
GREEN = "#a6e3a1"
BLUE = "#89b4fa"
RED = "#f38ba8"
YELLOW = "#f9e2af"


class App:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("TransPal — Trợ lý dịch thuật")
        self.root.geometry("800x580")
        self.root.configure(bg=BG)
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        self.ui_queue = queue.Queue()
        self.busy = threading.Lock()
        self.hotkey_handles = []
        self._loading_win = None

        self._build_ui()
        self._register_hotkeys()
        self.root.after(100, self._poll_queue)

        self.log(f"TransPal sẵn sàng. Phím tắt dịch: {CONFIG['hotkey']}")
        if CONFIG["mode"] == "ai" and not os.environ.get("GEMINI_API_KEY"):
            self.log("⚠ CẢNH BÁO: chưa đặt GEMINI_API_KEY → chế độ AI sẽ tự chuyển sang Google Translate.")
            self.log("   Cách đặt: mở CMD, gõ:  setx GEMINI_API_KEY \"key_của_bạn\"  rồi mở lại app.")

    # ---------- xây dựng giao diện ----------
    def _build_ui(self):
        top = tk.Frame(self.root, bg=BG)
        top.pack(fill="x", padx=12, pady=(10, 4))
        self.status_var = tk.StringVar()
        self._update_status()
        tk.Label(top, textvariable=self.status_var, bg=BG, fg=YELLOW,
                 font=("Segoe UI", 10, "bold")).pack(side="left")
        tk.Button(top, text="Đổi chế độ", command=self.toggle_mode,
                  bg=BLUE, fg="#11111b", relief="flat", padx=10,
                  cursor="hand2").pack(side="right")

        style = ttk.Style()
        try:
            style.theme_use("clam")
        except Exception:
            pass
        style.configure("TNotebook", background=BG, borderwidth=0)
        style.configure("TNotebook.Tab", padding=(14, 6))

        nb = ttk.Notebook(self.root)
        nb.pack(fill="both", expand=True, padx=12, pady=8)

        # --- TAB 1: NGỮ CẢNH ---
        tab_ctx = tk.Frame(nb, bg=BG)
        nb.add(tab_ctx, text="  Ngữ cảnh  ")

        btns = tk.Frame(tab_ctx, bg=BG)
        btns.pack(fill="x", pady=(8, 4), padx=8)
        self._btn(btns, "Import file (AI chuyển MD + merge)", self.import_file, GREEN).pack(side="left", padx=(0, 6))
        self._btn(btns, "Lưu ngữ cảnh", self.save_context, BLUE).pack(side="left", padx=(0, 6))
        self._btn(btns, "Tải lại từ file", self.reload_context, YELLOW).pack(side="left")
        tk.Label(btns, text="(AI đọc lại file này mỗi lần bạn bấm phím tắt dịch)",
                 bg=BG, fg=FG_DIM, font=("Segoe UI", 8)).pack(side="right")

        self.ctx_editor = scrolledtext.ScrolledText(
            tab_ctx, wrap="word", bg=BG2, fg=FG, insertbackground="white",
            font=("Consolas", 10), relief="flat", undo=True)
        self.ctx_editor.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        self.ctx_editor.insert("1.0", read_context_file())

        # --- TAB 2: CÀI ĐẶT ---
        tab_set = tk.Frame(nb, bg=BG)
        nb.add(tab_set, text="  Cài đặt  ")

        form = tk.Frame(tab_set, bg=BG)
        form.pack(anchor="nw", padx=16, pady=16)

        def row(label):
            f = tk.Frame(form, bg=BG)
            f.pack(fill="x", pady=5, anchor="w")
            tk.Label(f, text=label, bg=BG, fg=FG, width=24,
                     anchor="w", font=("Segoe UI", 10)).pack(side="left")
            return f

        r = row("Phím tắt dịch:")
        self.var_hotkey = tk.StringVar(value=CONFIG["hotkey"])
        tk.Entry(r, textvariable=self.var_hotkey, bg=BG2, fg=FG,
                 insertbackground="white", relief="flat", width=24).pack(side="left", ipady=3)

        r = row("Phím tắt đổi chế độ:")
        self.var_toggle = tk.StringVar(value=CONFIG["toggle_mode_hotkey"])
        tk.Entry(r, textvariable=self.var_toggle, bg=BG2, fg=FG,
                 insertbackground="white", relief="flat", width=24).pack(side="left", ipady=3)

        r = row("Chế độ dịch:")
        self.var_mode = tk.StringVar(value=CONFIG["mode"])
        ttk.Combobox(r, textvariable=self.var_mode, values=["ai", "basic"],
                     state="readonly", width=22).pack(side="left")

        r = row("Ngôn ngữ dịch sang:")
        self.var_lang = tk.StringVar(value=CONFIG["target_lang"])
        ttk.Combobox(r, textvariable=self.var_lang,
                     values=["vi", "en", "ja", "ko", "zh-CN", "fr", "de", "ru", "th"],
                     width=22).pack(side="left")

        r = row("Dán đè tự động:")
        self.var_auto = tk.BooleanVar(value=CONFIG["auto_replace"])
        tk.Checkbutton(r, variable=self.var_auto, bg=BG, activebackground=BG,
                       selectcolor=BG2,
                       text="Bỏ qua popup, thay thế ngay sau khi dịch xong",
                       fg=FG_DIM, font=("Segoe UI", 9)).pack(side="left")

        r = row("Khôi phục clipboard:")
        self.var_restore = tk.BooleanVar(value=CONFIG["restore_clipboard"])
        tk.Checkbutton(r, variable=self.var_restore, bg=BG, activebackground=BG,
                       selectcolor=BG2,
                       text="Trả lại nội dung clipboard cũ sau khi dán",
                       fg=FG_DIM, font=("Segoe UI", 9)).pack(side="left")

        r = row("Danh sách model AI:")
        self.txt_models = tk.Text(r, height=4, width=26, bg=BG2, fg=FG,
                                  insertbackground="white", relief="flat",
                                  font=("Consolas", 9))
        self.txt_models.pack(side="left")
        self.txt_models.insert("1.0", "\n".join(CONFIG["models"]))
        tk.Label(r, text="  (mỗi dòng một model, thử lần lượt từ trên xuống)",
                 bg=BG, fg=FG_DIM, font=("Segoe UI", 8)).pack(side="left")

        self._btn(form, "Lưu & áp dụng cài đặt", self.save_settings, GREEN).pack(anchor="w", pady=(16, 0))
        tk.Label(form, text="Ví dụ phím tắt hợp lệ: alt+shift+t, ctrl+alt+d, ctrl+shift+space",
                 bg=BG, fg=FG_DIM, font=("Segoe UI", 8)).pack(anchor="w", pady=(6, 0))

        # --- TAB 3: NHẬT KÝ ---
        tab_log = tk.Frame(nb, bg=BG)
        nb.add(tab_log, text="  Nhật ký  ")
        self.log_box = scrolledtext.ScrolledText(
            tab_log, wrap="word", bg="#11111b", fg=GREEN,
            font=("Consolas", 9), relief="flat", state="disabled")
        self.log_box.pack(fill="both", expand=True, padx=8, pady=8)

        tk.Label(self.root,
                 text="Cách dùng: bôi đen văn bản ở app bất kỳ → bấm phím tắt → xem/sửa bản dịch → Thay thế",
                 bg=BG, fg=FG_DIM, font=("Segoe UI", 9)).pack(pady=(0, 8))

    def _btn(self, parent, text, cmd, color):
        return tk.Button(parent, text=text, command=cmd, bg=color, fg="#11111b",
                         font=("Segoe UI", 9, "bold"), relief="flat",
                         padx=12, pady=4, cursor="hand2")

    def _update_status(self):
        mode_label = "AI (Gemini + ngữ cảnh)" if CONFIG["mode"] == "ai" else "Dịch nhanh (Google)"
        self.status_var.set(
            f"Chế độ: {mode_label}   |   Phím tắt: {CONFIG['hotkey']}   |   Dịch sang: {CONFIG['target_lang']}")

    # ---------- log & queue ----------
    def log(self, msg):
        def _append():
            self.log_box.configure(state="normal")
            self.log_box.insert("end", time.strftime("[%H:%M:%S] ") + str(msg) + "\n")
            self.log_box.see("end")
            self.log_box.configure(state="disabled")
        self.ui_queue.put(_append)

    def _poll_queue(self):
        try:
            while True:
                self.ui_queue.get_nowait()()
        except queue.Empty:
            pass
        self.root.after(100, self._poll_queue)

    # ---------- hotkey ----------
    def _register_hotkeys(self):
        for h in self.hotkey_handles:
            try:
                keyboard.remove_hotkey(h)
            except Exception:
                pass
        self.hotkey_handles = []
        try:
            self.hotkey_handles.append(
                keyboard.add_hotkey(CONFIG["hotkey"], self.on_translate_hotkey, suppress=True))
            self.hotkey_handles.append(
                keyboard.add_hotkey(CONFIG["toggle_mode_hotkey"],
                                    lambda: self.ui_queue.put(self.toggle_mode), suppress=True))
        except Exception as e:
            self.log(f"⚠ Không đăng ký được phím tắt: {e}")
            self.log("   Thử chạy lại bằng quyền Administrator nếu phím tắt không hoạt động.")

    def toggle_mode(self):
        CONFIG["mode"] = "basic" if CONFIG["mode"] == "ai" else "ai"
        self.var_mode.set(CONFIG["mode"])
        save_config(CONFIG)
        self._update_status()
        self.log(f"→ Đã chuyển chế độ dịch: {CONFIG['mode'].upper()}")

    def save_settings(self):
        CONFIG["hotkey"] = self.var_hotkey.get().strip() or CONFIG["hotkey"]
        CONFIG["toggle_mode_hotkey"] = self.var_toggle.get().strip() or CONFIG["toggle_mode_hotkey"]
        CONFIG["mode"] = self.var_mode.get()
        CONFIG["target_lang"] = self.var_lang.get().strip() or "vi"
        CONFIG["auto_replace"] = bool(self.var_auto.get())
        CONFIG["restore_clipboard"] = bool(self.var_restore.get())
        models = [m.strip() for m in self.txt_models.get("1.0", "end-1c").splitlines() if m.strip()]
        if models:
            CONFIG["models"] = models
        save_config(CONFIG)
        self._register_hotkeys()
        self._update_status()
        self.log("Đã lưu và áp dụng cài đặt mới.")
        messagebox.showinfo("TransPal", "Đã áp dụng cài đặt mới!")

    # ---------- ngữ cảnh ----------
    def save_context(self):
        write_context_file(self.ctx_editor.get("1.0", "end-1c"))
        self.log("Đã lưu context.md")

    def reload_context(self):
        self.ctx_editor.delete("1.0", "end")
        self.ctx_editor.insert("1.0", read_context_file())
        self.log("Đã tải lại context.md từ file.")

    def import_file(self):
        path = filedialog.askopenfilename(
            title="Chọn file ngữ cảnh / hướng dẫn dịch",
            filetypes=[("File hỗ trợ", "*.txt *.md *.docx *.pdf"),
                       ("Tất cả", "*.*")])
        if not path:
            return
        self.log(f"Import: {path}")
        threading.Thread(target=self._import_worker, args=(path,), daemon=True).start()

    def _import_worker(self, path):
        try:
            raw = extract_file_text(path)
            if not raw.strip():
                self.log("⚠ File rỗng hoặc không trích xuất được nội dung "
                         "(PDF dạng scan ảnh thì cần OCR).")
                return
            self.log(f"Đã đọc {len(raw)} ký tự. Đang nhờ AI chuyển sang Markdown và merge…")
            current = read_context_file()
            try:
                merged, model = ai_convert_and_merge(raw, current, log=self.log)
                self.log(f"AI ({model}) đã tạo bản merge. Đang hiện preview…")
            except Exception as e:
                self.log(f"⚠ AI không khả dụng ({e}). Sẽ nối nội dung thô vào cuối file.")
                merged = (current + "\n\n## Nội dung import từ "
                          + os.path.basename(path) + "\n" + raw.strip())
            self.ui_queue.put(lambda: self._show_merge_preview(merged))
        except Exception as e:
            self.log(f"⚠ Lỗi import: {e}")

    def _show_merge_preview(self, merged_text):
        win = tk.Toplevel(self.root)
        win.title("Xem trước — context.md sau khi merge")
        win.geometry("700x500")
        win.configure(bg=BG)
        win.attributes("-topmost", True)

        tk.Label(win, text="Đây là nội dung context.md MỚI sau khi AI merge. "
                           "Kiểm tra / sửa trực tiếp rồi bấm Áp dụng:",
                 bg=BG, fg=FG_DIM, font=("Segoe UI", 9),
                 wraplength=660, justify="left").pack(fill="x", padx=10, pady=(10, 4))

        box = scrolledtext.ScrolledText(win, wrap="word", bg=BG2, fg=FG,
                                        insertbackground="white",
                                        font=("Consolas", 10), relief="flat")
        box.pack(fill="both", expand=True, padx=10, pady=4)
        box.insert("1.0", merged_text)

        bf = tk.Frame(win, bg=BG)
        bf.pack(fill="x", padx=10, pady=(4, 10))

        def apply():
            final = box.get("1.0", "end-1c")
            write_context_file(final)
            self.ctx_editor.delete("1.0", "end")
            self.ctx_editor.insert("1.0", final)
            self.log("Đã áp dụng context.md mới (sau merge).")
            win.destroy()

        self._btn(bf, "Áp dụng (ghi vào context.md)", apply, GREEN).pack(side="left", padx=(0, 6))
        self._btn(bf, "Hủy", win.destroy, RED).pack(side="right")

    # ---------- hiệu ứng loading cạnh con trỏ ----------
    def _show_loading(self):
        if self._loading_win is not None:
            return
        win = tk.Toplevel(self.root)
        win.overrideredirect(True)          # không viền, không thanh tiêu đề
        win.attributes("-topmost", True)
        try:
            win.attributes("-alpha", 0.93)  # hơi trong suốt cho nhẹ nhàng
        except Exception:
            pass
        x, y = pyautogui.position()
        sw, sh = win.winfo_screenwidth(), win.winfo_screenheight()
        x = min(x + 18, sw - 160)
        y = min(y + 22, sh - 60)
        win.geometry(f"+{x}+{y}")
        lbl = tk.Label(win, text="⏳ Đang dịch", bg="#11111b", fg=YELLOW,
                       font=("Segoe UI", 10, "bold"), padx=14, pady=7)
        lbl.pack()
        self._loading_win = win
        self._loading_dots = 0

        def tick():
            if self._loading_win is not win:
                return  # đã bị đóng
            self._loading_dots = (self._loading_dots + 1) % 4
            try:
                lbl.config(text="⏳ Đang dịch" + "." * self._loading_dots)
                win.after(320, tick)
            except Exception:
                pass
        tick()

    def _hide_loading(self):
        win, self._loading_win = self._loading_win, None
        if win is not None:
            try:
                win.destroy()
            except Exception:
                pass

    # ---------- dịch bằng phím tắt ----------
    def on_translate_hotkey(self):
        if not self.busy.acquire(blocking=False):
            self.log("(đang xử lý yêu cầu trước, bỏ qua lần bấm này)")
            return
        threading.Thread(target=self._translate_worker, daemon=True).start()

    def _translate_worker(self):
        try:
            target_hwnd = get_foreground_hwnd()
            wait_hotkey_released(CONFIG["hotkey"])
            text, _old = get_selected_text()
            if not text.strip():
                self.log("⚠ Không tìm thấy văn bản nào đang được bôi đen.")
                return
            short = text if len(text) < 200 else text[:200] + " …"
            self.log(f"Gốc: {short}")
            self.ui_queue.put(self._show_loading)
            t0 = time.time()
            translated, source_label = do_translate(text, log=self.log)
            self.ui_queue.put(self._hide_loading)
            self.log(f"Dịch xong ({source_label}, {time.time()-t0:.1f}s)")
            if CONFIG.get("auto_replace"):
                paste_text(translated, target_hwnd)
            else:
                self.ui_queue.put(
                    lambda: self._show_translate_popup(text, translated, source_label, target_hwnd))
        except Exception as e:
            self.log(f"⚠ Lỗi bất ngờ: {e}")
        finally:
            self.ui_queue.put(self._hide_loading)
            self.busy.release()

    def _show_translate_popup(self, original, translated, source_label, target_hwnd):
        win = tk.Toplevel(self.root)
        win.title(f"TransPal — {source_label}")
        win.attributes("-topmost", True)
        x, y = pyautogui.position()
        w, h = 520, 360
        sw, sh = win.winfo_screenwidth(), win.winfo_screenheight()
        x = min(max(x - 60, 10), sw - w - 10)
        y = min(max(y + 20, 10), sh - h - 60)
        win.geometry(f"{w}x{h}+{x}+{y}")
        win.configure(bg=BG)

        tk.Label(win, text="Văn bản gốc:", fg=FG_DIM, bg=BG, anchor="w",
                 font=("Segoe UI", 9)).pack(fill="x", padx=10, pady=(8, 0))
        ob = scrolledtext.ScrolledText(win, height=4, wrap="word", bg=BG2, fg=FG,
                                       font=("Segoe UI", 10), relief="flat")
        ob.pack(fill="x", padx=10, pady=(2, 6))
        ob.insert("1.0", original)
        ob.configure(state="disabled")

        tk.Label(win, text="Bản dịch (sửa được trước khi thay thế):", fg=FG_DIM,
                 bg=BG, anchor="w", font=("Segoe UI", 9)).pack(fill="x", padx=10)
        tb = scrolledtext.ScrolledText(win, height=7, wrap="word", bg=BG2, fg=GREEN,
                                       font=("Segoe UI", 10),
                                       insertbackground="white", relief="flat")
        tb.pack(fill="both", expand=True, padx=10, pady=(2, 6))
        tb.insert("1.0", translated)

        bf = tk.Frame(win, bg=BG)
        bf.pack(fill="x", padx=10, pady=(0, 10))

        def on_replace(event=None):
            final = tb.get("1.0", "end-1c")
            win.destroy()
            threading.Thread(target=paste_text, args=(final, target_hwnd), daemon=True).start()

        def on_copy():
            pyperclip.copy(tb.get("1.0", "end-1c"))
            win.destroy()

        self._btn(bf, "Thay thế (Enter)", on_replace, GREEN).pack(side="left", padx=(0, 6))
        self._btn(bf, "Chỉ copy", on_copy, BLUE).pack(side="left", padx=(0, 6))
        self._btn(bf, "Đóng (Esc)", win.destroy, RED).pack(side="right")

        win.bind("<Return>", on_replace)
        win.bind("<Escape>", lambda e: win.destroy())
        win.focus_force()
        tb.focus_set()

    # ---------- thoát ----------
    def on_close(self):
        try:
            keyboard.unhook_all_hotkeys()
        except Exception:
            pass
        self.root.destroy()

    def run(self):
        self.root.mainloop()


def main():
    try:
        read_context_file()
        App().run()
    except Exception:
        # Chạy bằng .pyw không có console → hiện lỗi bằng hộp thoại
        err = traceback.format_exc()
        try:
            r = tk.Tk(); r.withdraw()
            messagebox.showerror("TransPal — lỗi khởi động", err[-1500:])
        except Exception:
            pass


if __name__ == "__main__":
    main()
